from flask import Flask, render_template, request, redirect, url_for, send_from_directory, flash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
import os
from werkzeug.utils import secure_filename
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from base64 import urlsafe_b64decode, urlsafe_b64encode
from datetime import datetime, timedelta
import re
import time
import pandas as pd
from tabulate import tabulate
import logging
import io
from docx import Document
from PyPDF2 import PdfReader
import shutil

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'Resumes'
app.config['SECRET_KEY'] = 'your-secret-key-here'  # Change to a secure random key

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Flask-Login setup
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']

# User class for Flask-Login
class User(UserMixin):
    def __init__(self, id):
        self.id = id

# User store with your provided credentials
users = {
    "maneeshaupender30@gmail.com": {"password": "Chawo30o@"},
    "saicharan.rajampeta@iitlabs.us": {"password": "Db2@Admin"}
}

@login_manager.user_loader
def load_user(user_id):
    if user_id in users:
        return User(user_id)
    return None

# --- All Original Functions ---
def get_emails_by_job_id(service, job_id):
    """Fetch all emails containing a given job ID within the last 30 days."""
    today = datetime.now()
    last_month = today - timedelta(days=30)
    last_month_str = last_month.strftime('%Y/%m/%d')
    all_messages = []
    page_token = None
    while True:
        try:
            results = service.users().messages().list(
                userId="me", q=f"after:{last_month_str} {job_id}", maxResults=500, pageToken=page_token
            ).execute()
            messages = results.get("messages", [])
            all_messages.extend(messages)
            page_token = results.get("nextPageToken")
            if not page_token:
                break
            time.sleep(1)
        except Exception as e:
            logging.error(f"Error fetching emails: {e}")
            break
    return all_messages

def decode_base64(data):
    """Decodes base64 email content safely."""
    missing_padding = len(data) % 4
    if missing_padding:
        data += '=' * (4 - missing_padding)
    return urlsafe_b64decode(data)

def extract_email_body(payload):
    """Extracts the email body content."""
    if not payload:
        return ""
    if "body" in payload and "data" in payload["body"]:
        return decode_base64(payload["body"]["data"]).decode("utf-8", errors="ignore")
    if "parts" in payload:
        for part in payload["parts"]:
            if part.get("mimeType", "") in ["text/plain", "text/html"] and "data" in part.get("body", {}):
                return decode_base64(part["body"]["data"]).decode("utf-8", errors="ignore")
            if "parts" in part:
                nested_body = extract_email_body(part)
                if nested_body:
                    return nested_body
    return ""

def extract_skill_names(skills):
    """Extracts skill names from skill sentences and preserves skill type."""
    skill_names = []
    for skill in skills:
        skill_type_match = re.search(r"(Required|Highly desired|Nice to have)\s+\d+\s+Months", skill, re.IGNORECASE)
        skill_type = skill_type_match.group(0) if skill_type_match else ""
        skill_name_match = re.search(
            r"(?:Ability\s+to\s+|Understanding\s+|Strong\s+knowledge\s+of\s+|Experience\s+or\s+interest\s+in\s+|Ability\s+to\s+learn\s+|Excellent\s+|Knowledge\s+of\s+)(.*?)\s*(?:Required|Highly desired|Nice to have|\d+\s+Months|$)",
            skill, re.IGNORECASE
        )
        if skill_name_match:
            skill_name = skill_name_match.group(1).strip()
            skill_name = skill_name[0].upper() + skill_name[1:]
            skill_names.append(f"• {skill_name} {skill_type}")
        else:
            skill_names.append(f"• {skill}")
    return skill_names

def extract_skills(email_body):
    """Extracts Skills section while preserving 'Required X Years' text."""
    match = re.search(r"Skills?:\s*(.*?)\s*(?:Responsibilities:|Qualifications:|Description:|Job ID:|$)", email_body, re.DOTALL | re.IGNORECASE)
    if not match:
        return None
    skill_text = match.group(1).strip()
    return clean_skill_text(skill_text)

def clean_skill_text(skill_text):
    """Cleans formatting while retaining skill text with years."""
    if not skill_text:
        return None
    skill_text = re.sub(r"(\n\s*[-•*]\s*\d+\.?\s*|\n\s*[-•*]\s*)", "\n", skill_text)
    skill_text = re.sub(r"^\d+\.\s*", "", skill_text, flags=re.MULTILINE)
    skills = []
    current_skill = ""
    for line in skill_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if re.match(r"^(Ability|Understanding|Strong|Experience|Knowledge|Excellent)", line, re.IGNORECASE):
            if current_skill:
                skills.append(current_skill.strip())
            current_skill = line
        else:
            current_skill += " " + line
    if current_skill:
        skills.append(current_skill.strip())
    return extract_skill_names(skills)

def extract_details_from_body(body):
    """Extracts comprehensive candidate details from email body."""
    name_patterns = [
        r"First Name\s*\(.*?\):\s*(.*?)\s*Middle Name\s*\(.*?\):\s*(.*?)\s*Last Name\s*\(.*?\):\s*(.*?)(?:\n|$)",
        r"Name\s*:\s*(.*?)(?:\n|$)",
        r"<(b|strong)>(.*?)</\1>"
    ]
    name = "N/A"
    for pattern in name_patterns:
        match = re.search(pattern, body, re.IGNORECASE | re.DOTALL)
        if match:
            groups = [g.strip() for g in match.groups() if g]
            cleaned_groups = [re.sub(r'\s*>\s*', ' ', g) for g in groups]
            name = " ".join(filter(None, cleaned_groups))
            break
    phone_match = re.search(r"(?i)(?:phone\s*#|phone\s*no|phone|ph|mobile|contact)[\s#:]*([+]?[\d]{0,4}[\s\-.]*\(?\d{0,4}\)?[\s\-.]*\d{3,4}[\s\-.]*\d{3,4})", body)
    phone = phone_match.group(1).strip() if phone_match else "N/A"
    email_match = re.search(r"(?i)email\s*[:\-]?\s*([\w\.-]+@[\w\.-]+\.\w+)", body)
    email = email_match.group(1).strip() if email_match else "N/A"
    location_match = re.search(r"(?i)Current location\s*(?:\(city/state\))?\s*[:\-]?\s*(?:[\*\b_]{1,2})?\s*([\w\s.,-]+(?:[/,]\s*[\w\s.-]+)?)\s*(?:[\*\b_]{1,2})?\s*(?=\n|$)", body)
    current_location = location_match.group(1).strip() if location_match else "N/A"
    experience = "0 years"
    experience_patterns = [
        r"Total no of years experience:\s*([\d]+(?:\.\d+)?\+?)",
        r"Years of experience:\s*([\d]+(?:\.\d+)?\+?)",
        r"Experience:\s*([\d]+(?:\.\d+)?\+?)",
        r"Experience:\s*([\d.+]*)"
    ]
    for pattern in experience_patterns:
        match = re.search(pattern, body, re.IGNORECASE)
        if match:
            exp_value = match.group(1).strip()
            experience = f"{exp_value} years" if not exp_value.lower().endswith("years") else exp_value
            break
    certification_count = 0
    cert_match = re.search(r"Certification Count:\s*(\d+)", body, re.IGNORECASE)
    if cert_match:
        certification_count = int(cert_match.group(1).strip())
    govt_exp_match = re.search(r"(?i)Government\s*experience\s*:\s*(?:\(mention\s*the\s*government\s*name's\s*in\s*resume\s*otherwise\s*No\)\s*:?\s*)?([\w\s,&.\-]+?(?:&\w+;)?[\w\s,&.\-]*?)(?:\n|$|\r\n)", body, re.IGNORECASE)
    if govt_exp_match:
        government_experience = govt_exp_match.group(1).strip()
        government_experience = re.sub(r'\s+', ' ', government_experience)
        if government_experience.lower() in ["no", "not mentioned", "none"]:
            government_experience = "Not worked with the government"
    else:
        government_experience = "Not worked with the government"
    visa_info = "N/A"
    visa_patterns = [
        r"Visa\s*Status\s*with\s*Validity\s*:\s*([^\n\r]*)",
        r"Visa\s*type\s*and\s*sponsor\s*name\s*\(.*?\)\s*:\s*([^\n\r]*)",
        r"Visa\s*type\s*:\s*([^\n\r]*)",
        r"Status\s*:\s*([^\n\r]*)"
    ]
    for pattern in visa_patterns:
        match = re.search(pattern, body, re.IGNORECASE)
        if match:
            visa_info = match.group(1).strip()
            if not visa_info or visa_info.isspace():
                visa_info = "N/A"
            break
    return {
        "Candidate Name": name, "Phone No": phone, "Email": email, "Current Location": current_location,
        "Total Experience": experience, "Certification Count": certification_count,
        "Government Experience": government_experience, "Visa Status": visa_info.strip()
    }

def extract_text_from_attachment(attachment_data, filename):
    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(attachment_data))
            text = "\n".join(page.extract_text() for page in reader.pages)
            return text
        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(attachment_data))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        else:
            logging.warning(f"Unsupported file type: {filename}")
            return None
    except Exception as e:
        logging.error(f"Error extracting text from {filename}: {e}")
        return None

def process_attachments(service, message_id):
    """Collects valid (filename, attachment_id) tuples."""
    try:
        msg = service.users().messages().get(userId="me", id=message_id, format="full").execute()
        payload = msg.get("payload", {})
        parts = payload.get("parts", [])
        attachments = []
        for part in parts:
            filename = part.get("filename")
            attachment_id = part.get("body", {}).get("attachmentId")
            if filename and attachment_id:
                attachments.append((filename, attachment_id))
        return attachments
    except Exception as e:
        logging.error(f"Error listing attachments: {e}")
        return []

def get_attachment_data(service, message_id, attachment_id):
    """Fetches and decodes attachment data using the Gmail API."""
    try:
        attachment = service.users().messages().attachments().get(
            userId="me", messageId=message_id, id=attachment_id
        ).execute()
        data = attachment.get("data", "")
        return urlsafe_b64decode(data)
    except Exception as e:
        logging.error(f"Error fetching attachment data: {e}")
        return None

def filter_excluded_files(attachments):
    exclusion_terms = ["rtr", "sow", "sm", "jd", "job", "description", "mail", "contract", "project", "h1", "gc", "dl", "signed"]
    filtered = []
    for filename in attachments:
        lower_name = filename.lower()
        if any(term in lower_name for term in exclusion_terms):
            continue
        if filename.lower().endswith((".pdf", ".docx")):
            filtered.append(filename)
    return filtered

def is_resume_content(text):
    required_sections = [
        r"work\s*experience", r"education", r"skills", r"summary", r"projects", r"certifications",
        r"expertise\s*brief", r"professional\s*engagements", r"academic\s*qualifications",
        r"certificates\s*&\s*accolades", r"key\s*skills", r"executive\s*summary", r"technical\s*skills",
        r"professional\s*summary", r"career\s*highlights", r"professional\s*milestones",
        r"responsibilities", r"qualification\s*badges", r"technical\s*profile",
        r"significant\s*practices", r"certification", r"technical\s*expertise", r"client",
        r"education\s*&\s*credentials", r"business\s*process\s*improvement", r"career\s*achievements",
        r"summary\s*of\s*the\s*experience", r"work\s*/\s*assignment\s*history",
        r"education\s*and\s*professional\s*qualifications", r"work\s*/\s*assignment\s*history"
    ]
    found = sum(1 for section in required_sections if re.search(section, text, re.IGNORECASE))
    return found >= 1

def identify_resume(service, message_id, attachments):
    """Identifies the resume from pre-validated attachments."""
    valid_files = filter_excluded_files([fn for (fn, _) in attachments])
    for filename in valid_files:
        attachment_id = next((aid for (fn, aid) in attachments if fn == filename), None)
        if not attachment_id:
            continue
        attachment_data = get_attachment_data(service, message_id, attachment_id)
        if not attachment_data:
            continue
        text = extract_text_from_attachment(attachment_data, filename)
        if text and is_resume_content(text):
            return filename
    return "N/A"

def extract_skills_from_subject(subject):
    """Extracts skills dynamically from the email subject, ignoring any prefix."""
    if isinstance(subject, list):
        return subject
    skill_pattern = r"(?:.*\bwith\b\s*|\bincluding\b\s*)(.*)"
    match = re.search(skill_pattern, subject, re.IGNORECASE)
    if match:
        skills_part = match.group(1).strip()
        skills_list = [skill.strip() for skill in skills_part.split(",")]
        return skills_list
    else:
        return [subject]

def extract_email_data(service, message_id):
    """Fetches the email and extracts details, skills, and subject."""
    try:
        msg = service.users().messages().get(userId="me", id=message_id, format="full").execute()
        headers = msg.get("payload", {}).get("headers", [])
        subject = next((header["value"] for header in headers if header["name"] == "Subject"), "N/A")
        subject_skills = extract_skills_from_subject(subject)
        email_body = extract_email_body(msg.get("payload", {}))
        if not email_body:
            return None, None, None
        email_skills = extract_skills(email_body)
        details = extract_details_from_body(email_body)
        attachment_results = process_attachments(service, message_id)
        filenames = [fn for (fn, _) in attachment_results]
        resume_filename = identify_resume(service, message_id, attachment_results)
        details["Resume File"] = resume_filename
        return details, email_skills, subject_skills
    except Exception as e:
        logging.error(f"Error processing email {message_id}: {e}")
        return None, None, None

def create_resume_folder(folder_name="Resumes"):
    if os.path.exists(folder_name):
        for filename in os.listdir(folder_name):
            file_path = os.path.join(folder_name, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                logging.error(f"Failed to delete {file_path}: {e}")
        logging.info(f"Cleared existing files in folder: {folder_name}")
    else:
        os.makedirs(folder_name)
        logging.info(f"Created folder: {folder_name}")
    return folder_name

def save_resumes_to_folder(service, details, message_id, attachments, resume_folder):
    resume_filename = details["Resume File"]
    if resume_filename == "N/A":
        return
    attachment_id = next((aid for (fn, aid) in attachments if fn == resume_filename), None)
    if not attachment_id:
        return
    attachment_data = get_attachment_data(service, message_id, attachment_id)
    if not attachment_data:
        return
    resume_file_path = os.path.join(resume_folder, resume_filename)
    try:
        with open(resume_file_path, "wb") as file:
            file.write(attachment_data)
    except Exception as e:
        logging.error(f"Failed to save resume file {resume_filename}: {e}")

def extract(path):
    if path.endswith(".docx"):
        doc = Document(path)
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        return full_text
    elif path.endswith(".pdf"):
        reader = PdfReader(path)
        full_text = ""
        for page in reader.pages:
            full_text += page.extract_text() + "\n"
        return full_text
    else:
        raise ValueError("Unsupported file type. Only DOCX and PDF files are supported.")

def match(full_text, subject_skills):
    text_lower = full_text.lower()
    skill_counts = {}
    for skill in subject_skills:
        pattern = r'\b' + re.escape(skill.strip().lower()) + r'\b'
        matches = re.findall(pattern, text_lower)
        count = len(matches)
        if count > 0:
            skill_counts[skill.strip()] = count
    return skill_counts

def compare_skills_with_resumes(subject_skills, resume_folder):
    matched_skills_dict = {}
    for filename in os.listdir(resume_folder):
        if filename.endswith((".docx", ".pdf")):
            file_path = os.path.join(resume_folder, filename)
            try:
                full_text = extract(file_path)
                skill_counts = match(full_text, subject_skills)
                if skill_counts:
                    matched_skills_dict[filename] = skill_counts
            except Exception as e:
                print(f"Error processing {filename}: {e}")
    return matched_skills_dict

def calculate_resume_score(matched_skills, government_experience):
    score = 0
    if matched_skills and isinstance(matched_skills, str):
        matched_skills_list = [skill.strip() for skill in matched_skills.split(",")]
        score += len(matched_skills_list) * 5
    if government_experience.lower() != "not worked with the government":
        government_count = government_experience.count(",") + 1
        score += government_count * 10
    return score

# --- Flask Routes ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')  # Matches your login.html
        password = request.form.get('password')
        if username in users and users[username]['password'] == password:
            user = User(username)
            login_user(user)
            return redirect(url_for('index'))
        else:
            flash('Invalid username or password', 'danger')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    if request.method == 'POST':
        job_id = request.form.get('job_id')
        if job_id:
            return redirect(url_for('process_job', job_id=job_id))
    return render_template('index.html')

@app.route('/process/<job_id>')
@login_required
def process_job(job_id):
    try:
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
        service = build("gmail", "v1", credentials=creds)
    except Exception as e:
        logging.error(f"Authentication failed: {e}")
        return render_template('error.html', error="Authentication failed. Please check your credentials.")
    
    messages = get_emails_by_job_id(service, job_id)
    if not messages:
        return render_template('results.html', job_id=job_id, message="No emails found related to this job ID.")
    
    resume_folder = create_resume_folder(app.config['UPLOAD_FOLDER'])
    subject_skills_list = []
    email_data = []
    
    for msg in messages:
        message_id = msg["id"]
        details, skills, subject_skills = extract_email_data(service, message_id)
        if skills:
            subject_skills_list.append({"Subject": subject_skills, "Skills": skills})
        if details:
            email_data.append(details)
            attachments = process_attachments(service, message_id)
            if attachments:
                save_resumes_to_folder(service, details, message_id, attachments, resume_folder)
    
    if subject_skills_list and email_data:
        df = pd.DataFrame(email_data)
        subject_skills = subject_skills_list[0]["Subject"]
        matched_skills_dict = compare_skills_with_resumes(subject_skills, resume_folder)
        
        matched_skills_list = []
        for index, row in df.iterrows():
            resume_filename = row["Resume File"]
            if resume_filename in matched_skills_dict:
                skills_str = ", ".join([f"{skill} ({count})" for skill, count in matched_skills_dict[resume_filename].items()])
                matched_skills_list.append(skills_str)
            else:
                matched_skills_list.append("N/A")
        
        df["Matched Skills"] = matched_skills_list
        resume_scores = [calculate_resume_score(row["Matched Skills"], row["Government Experience"]) for index, row in df.iterrows()]
        df["Resume Score"] = resume_scores
        df["Rank"] = df["Resume Score"].rank(ascending=False, method="min").astype(int)
        df = df.sort_values(by="Rank")
        columns_order = ["Rank"] + [col for col in df.columns if col != "Rank"]
        df = df[columns_order]
        
        table_html = df.to_html(classes='table table-striped', index=False)
        skills_info = [{"subject": item['Subject'], 'skills': item['Skills']} for item in subject_skills_list]
        
        return render_template('results.html', job_id=job_id, table_html=table_html, skills_info=skills_info, num_candidates=len(df))
    else:
        return render_template('results.html', job_id=job_id, message="No skills or resumes found for comparison.")

@app.route('/download/<filename>')
@login_required
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

if __name__ == '__main__':
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run(debug=True)
